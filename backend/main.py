from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Inches
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from flask_cors import CORS

app = Flask(__name__)
CORS(app) 

import io
from docx import Document
from zipfile import ZipFile, ZIP_DEFLATED

def parse_ooxml_in_memory(ooxml_content):
    """直接在内存中解析OOXML字符串并提取标题与段落"""
    try:
        # 将字符串转换为BytesIO对象，模拟文件
        ooxml_bytes = io.BytesIO(ooxml_content.encode('utf-8'))
        
        # 使用BytesIO创建ZipFile对象
        with ZipFile(ooxml_bytes, mode='r') as zipf:
            # 检查Word文档的基本结构是否存在，比如document.xml
            if 'word/document.xml' not in zipf.namelist():
                print("The provided content does not seem to contain a valid Word document structure.")
                return [], []
            
            # 读取document.xml内容
            document_xml = zipf.read('word/document.xml').decode('utf-8')
            
            # 将document.xml内容写入一个新的BytesIO，以便用docx处理
            docx_io = io.BytesIO()
            with ZipFile(docx_io, 'w', compression=ZIP_DEFLATED) as docx_zip:
                # 添加必要的目录结构和文件到新的ZipFile
                docx_zip.writestr('word/document.xml', document_xml)
                # 其他必要文件如'[Content_Types].xml'等也需手动添加，这里简化处理
            
            # 重置BytesIO的读取指针
            docx_io.seek(0)
            
            # 使用Document处理内存中的ZipFile
            doc = Document(docx_io)
            
            headings = [(para.style.name, para.text) for para in doc.paragraphs if para.style.name.startswith('Heading')]
            paragraphs = [para.text for para in doc.paragraphs if not para.style.name.startswith('Heading')]
            
            return headings, paragraphs
    
    except Exception as e:
        print(f"Error during processing: {e}")
        return [], []


@app.route('/ooxml', methods=['POST'])
def receive_and_create_word():
    try:
        # 获取请求中的JSON数据
        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({"error": "缺少必要的'content'字段"}), 400

        # 提取OOXML内容
        ooxml_content = data['content']

        # 创建一个空的Word文档对象
        # doc = Document()

        # 解析OOXML数据并添加到文档中
        # 注意：此处简化处理，实际应用中需要根据OOXML的具体结构进行解析
        # 假设ooxml_content是一个完整的Word文档的XML表示，我们直接尝试构建新的文档
        # 实际情况可能需要更复杂的处理逻辑来正确解析并合并不同的XML元素
        # doc = parse_xml(ooxml_content)

        headings, paragraphs =  parse_ooxml_in_memory(ooxml_content)
        if headings:
            print("Headings:")
            for level, text in headings:
                print(f"{level}: {text}")

        if paragraphs:
            print("\nParagraphs:")
            for para in paragraphs:
                print(para)
        
        
        doc = Document(docx=doc)
        # 尝试将解析后的XML直接附加到文档的body中
        # 初始化存储标题和段落的列表
        
        headings = []
        paragraphs = []
        
        # 遍历文档中的段落
        for para in doc.paragraphs:
            # 检查段落是否有样式，判断是否为标题
            style_name = para.style.name
            if 'Heading' in style_name:  # 假设标题样式以'Heading'开头
                headings.append((style_name, para.text))
            else:
                paragraphs.append(para.text)

        output_path = 'data/received_content.docx'
        return jsonify({"message": f"Word文档已成功从OOXML数据创建并保存至{output_path}"}), 200
        
        # 打印提取的信息
        print("Headings:")
        for level, text in headings:
            print(f"Level {level.split(' ')[1]}: {text}")
            
        print("\nParagraphs:")
        for para_text in paragraphs:
            print(para_text)
        # 保存文档到服务器（示例路径，实际部署时应考虑合适的存储策略和安全问题）
        output_path = 'data/received_content.docx'
        return jsonify({"message": f"Word文档已成功从OOXML数据创建并保存至{output_path}"}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, port=8000)